#Importing Libraries
import csv
import tkinter as tk
from tkinter import messagebox as msg
from tkinter import messagebox 
from tkinter import StringVar
from tkinter import OptionMenu
from tkinter import Radiobutton
from tkinter import IntVar
from tkinter import PhotoImage
from tkinter import Canvas
from tkinter import Spinbox
from datetime import date
import socket
import sys
from tkinter import ttk
import random
import boto3
from tkinter.simpledialog import askstring
import mysql.connector
import os
import ssl
from PIL import ImageTk, Image
import webbrowser
import numpy as np
import matplotlib.pyplot as plt
import smtplib
from datetime import date
import folium
from geopy.geocoders import Nominatim
import urllib.request
import json
import smtplib
import config
import docx
from math import radians, cos, sin, asin, sqrt
import Func
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from Calendar import Calendar
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from CalendarView import CalendarView

#Connecting to the database
mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="",
    database="maindatabase"
)

#Connecting to AWS Services for tests
client = boto3.client(
    "sns",
    aws_access_key_id="",
    aws_secret_access_key="",
    region_name="eu-west-1"
)


print(mydb)

mycursor= mydb.cursor()
#Creating all the tables if they're not already created.
mycursor.execute("""CREATE TABLE IF NOT EXISTS manager(managerID int,managerUsername VARCHAR(255),managerPassword VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS electrician(electricianID int,electricianUsername VARCHAR(255),electricianPassword VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS clientinfo (clientID VARCHAR(255),firstname VARCHAR(255),secondname VARCHAR(255),
age VARCHAR(255),gender VARCHAR(255),address VARCHAR(255),postcode VARCHAR(255),userID VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS usercredentials (userID VARCHAR(255),username VARCHAR(255),password VARCHAR(255),stakeholder VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS electricianinfo(electricianID int,electricianName VARCHAR(255),electricianSurname VARCHAR(255),electricianAddress VARCHAR(255),electricianHourlyWage VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS activeJobs(electricianID VARCHAR(255),firstname VARCHAR(255),secondname VARCHAR(255),gender VARCHAR(255),jobtype VARCHAR(255),hourlywage VARCHAR(255),PostcodeOfWork VARCHAR(255),jobStatus VARCHAR(45))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS clientContact(clientID int,clientPhoneNumber VARCHAR(255),clientEmailAddress VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS managerContact(managerID int,managerPhoneNumber VARCHAR(255),managerEmailAddress VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS currentStock(stockType VARCHAR(255),currentAmount VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS clientFeedback(SystemFeedback VARCHAR(255),timeFeedback VARCHAR(255),FriendlinessFeedback VARCHAR(255),CostFeedback VARCHAR(255),OverallFeedback VARCHAR(255))""")
mycursor.execute("""CREATE TABLE IF NOT EXISTS pastJobs(electricianID VARCHAR(255),firstname VARCHAR(255),secondname VARCHAR(255),gender VARCHAR(255),jobtype VARCHAR(255),hourlywage VARCHAR(255),PostcodeOfWork VARCHAR(255),jobStatus VARCHAR(45))""")


#Main class
class Main(tk.Tk):
    def __init__(self, *args, **kwargs):
        tk.Tk.__init__(self, *args, **kwargs)
        container = tk.Frame(self)
        # Window siz
        tk.Tk.geometry(self, "600x600")
        # Window title
        tk.Tk.wm_title(self, "Darson Services")

        container.pack(side="top", fill="both", expand=True)
        # Setting up rows and columns for the base container
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)
        # Allows the frames to be created from the container
        self.frames = {}
        # For loop - looping through all the frames
        for F in (LoginFrame,MainWindow,ManagerWindow,ResetPassword,BookElectrician,CustomerFeedback,ViewStatusOfJobs,ElectricianManageStock,ManagerCreateNewAccount2,PostCodeLookup,ElectricianWindow,ElectricianActiveJobs,ManagerReports,ViewCurrentJobs,ManagerCreateNewAccount,ManagerAccouncementFrame,ManagerManageStock,ViewAccounts):
            frame = F(container, self)

            self.frames[F] = frame

            frame.grid(row=0, column=0, sticky="news")
        # First frame that will be shown
        self.show_frame(LoginFrame)

    def show_frame(self, cont):
        frame = self.frames[cont]
        # Running the tk raise function
        frame.tkraise()

    def Clear(self, Username, Password):
          Username.set("")
          Password.set("")



    def Check_Login(self,Username,Password,LoginType):
        #Getting the Username and Password
        LoginUsername = Username.get()
        LoginPassword = Password.get()
        CheckLoginType = LoginType.get()
        print(CheckLoginType)
        print(LoginUsername)
        print(LoginPassword)
        #Presence Check for both Username and Password
        if (LoginPassword == ""):
            msg.showinfo("Error","Invalid Username")
            return
        if (LoginUsername == ""):
            msg.showinfo("Error","Invalid Password")
            return
        LoginFound = False
        #Checking Logintype
        if CheckLoginType == "Customer":
            #Getting data from SQL DB to compare
            mycursor.execute("SELECT userID,username,password,stakeholder FROM usercredentials")
            myresult = mycursor.fetchall()
            for row in myresult:
                print("Checking 2")
                #Bypassing a blank row
                if row == []:
                    pass
                #Comparing the entered username and password to the ones in the database
                if (str(row[1]) == str(LoginPassword) and str(row[2]) == str(LoginUsername)):
                    self.show_frame(MainWindow)
                    #Boolean set to true
                    LoginFound = True
                    #Message box to display it worked
                    msg.showinfo("Successful", "Successfully logged in")
                    #Resetting Username and Password
                    Username.set("")
                    Password.set("")
                    #Setting the variables to blank
                    LoginUsername = ""
                    LoginPassword = ""
                    break

                    
            if (LoginFound == False):
                #Outputting error message to the user.
                msg.showinfo("Error", "Invalid Login details")
    
        #Same Login type check as above, along with the same proccess, described in the annotation
        elif CheckLoginType == "Electrician":
            mycursor.execute("SELECT userID,username,password,stakeholder FROM usercredentials")
            myresult = mycursor.fetchall()
            for row in myresult:
                    if row == []:
                        pass
                    if (str(row[1]) == str(LoginPassword) and str(row[2]) == str(LoginUsername)):
                        self.show_frame(ElectricianWindow)
                        LoginFound = True
                        msg.showinfo("Successful", "Successfully logged in")
                        Username.set("")
                        Password.set("")
                        LoginUsername = ""
                        break
                        LoginPassword = ""
                        csvfile.close()
            if LoginFound == False:
                msg.showinfo("Error", "Invalid Login details")
        #Same Login type check as above, along with the same proccess, described in the annotation
        elif CheckLoginType == "Manager":
            mycursor.execute("SELECT userID,username,password,stakeholder FROM usercredentials")
            myresult = mycursor.fetchall()
            for row in myresult:
                    print(row[0])
                    print(row[1])
                    print(row[2])
                    if row == []:
                        pass
                    if (str(row[1]) == str(LoginPassword) and str(row[2]) == str(LoginUsername)):
                        self.show_frame(ManagerWindow)
                        LoginFound = True
                        msg.showinfo("Successful", "Successfully logged in")
                        LoginUsername = ""
                        LoginPassword = ""
                        Username.set("")
                        Password.set("")
                        break
                    
            if LoginFound == False:
                msg.showinfo("Error", "Invalid Login details")
        #Presence check
        elif CheckLoginType == "":
            msg.showinfo("Error","Failed to complete all fields")

    #Back button function
    def MainMenuBack(self):
        self.show_frame(LoginFrame)
    #Reset password function, allows the user to reset their password
    def ResetPassword(self,ResetUsername, ResetLoginType, ResetPhoneNumber):
        #Getting the Username, PhoneNumber and LoginType
        ResetLoginUsername1 = ResetUsername.get()
        ResetLoginType1 = ResetLoginType.get()
        ResetPhoneNumber1 = ResetPhoneNumber.get()
        print(ResetLoginType1)
        print(ResetPhoneNumber1)
        Reset = False
        #Format Check
        if (str(ResetPhoneNumber1[0]) != "+"):
            #Format Check failed, asking them to try again.
            msg.showinfo("Error","Please input your Phone Number in the correct format, starting with a +")
            return
        #Used for debugging purpsoes, displaying that 1st check is complete.
        print("Validation Check 1 Passed")
        #Length Check on phone number
        if(len(ResetPhoneNumber1) != 13):
            #Unsuccessful message
            msg.showinfo("Error","Invalid Phone Number please try again")
            return
        #Used for debugging purpsoes, displaying that 2nd check is complete.
        print("Validation Check 2 Passed")
        #Presence Check
        if(ResetLoginUsername1 == ""):
            msg.showinfo("Error","Invalid Username please try again")
            return
        print("Validation Check 3 Passed")
        #Checking the user group
        mycursor.execute("SELECT userID,username,password,stakeholder FROM usercredentials")
        myresult = mycursor.fetchall()
        for row in myresult:
            if row[1] == ResetLoginUsername1 and row[3] == ResetLoginType1:
                if (ResetLoginType1 == "Manager"):
                    print(ResetPhoneNumber1)
                    print("Check1")
                    #Generating Recovery Code
                    RecoveryCode = random.randint(1000,9999)
                    print(RecoveryCode)
                    #Sending text message to the number, with the recovery code.
                    client.publish(
                        PhoneNumber=str(ResetPhoneNumber1),
                        Message=str(RecoveryCode)
                        )
                    print("Check2")
                    #Asking them to input the recovery code
                    ResetCodeCheck = askstring("Enter the Reset Code", "Please enter your Reset Code sent to your phone number:")
                    #Checking if the code they entered and the one orginally generated match.
                    if (int(ResetCodeCheck) == RecoveryCode):
                        Reset = True
                        #Asking them to input their new password.
                        NewPassword = askstring("Success", "Enter your new password")
                        mycursor = mydb.cursor()
                        #updating the password of a specific account
                        UpdateSQL = "UPDATE usercredentials SET Password = %s WHERE Username = %s"
                        passworddata = (NewPassword), (ResetLoginUsername1)
                        mycursor.execute(UpdateSQL, passworddata)
                        mydb.commit()
                        self.show_frame(LoginFrame)
                    else:
                        #Displaying to the user they inputted an invalid reset code.
                        msg.showinfo("Error","Invalid Reset Code")
                        return

                #Same Operation as above, just with the different user group
                if (ResetLoginType1 == "Electrician"):
                    print(ResetPhoneNumber1)
                    print("Check1")
                    RecoveryCode = random.randint(1000,9999)
                    print(RecoveryCode)
                    client.publish(
                        PhoneNumber=str(ResetPhoneNumber1),
                        Message=str(RecoveryCode)
                        )
                    print("Check2")
                    ResetCodeCheck = askstring("Enter the Reset Code", "Please enter your Reset Code sent to your phone number:")
                    if (int(ResetCodeCheck) == RecoveryCode):
                        Reset = True
                        NewPassword = askstring("Success", "Enter your new password")
                        mycursor = mydb.cursor()
                        #updating the password of a specific account
                        UpdateSQL = "UPDATE usercredentials SET Password = %s WHERE Username = %s"
                        passworddata = (NewPassword), (ResetLoginUsername1)
                        mycursor.execute(UpdateSQL, passworddata)
                        mydb.commit()
                        self.show_frame(LoginFrame)
                    else:
                        msg.showinfo("Error","Invalid Reset Code")
                        return

                #Same Operation as above, just with the different user group
                if (ResetLoginType1 == "Customer"):
                    print(ResetPhoneNumber1)
                    print("Check1")
                    RecoveryCode = random.randint(1000,9999)
                    print(RecoveryCode)
                    client.publish(
                        PhoneNumber=str(ResetPhoneNumber1),
                        Message=str(RecoveryCode)
                        )
                    print("Check2")
                    ResetCodeCheck = askstring("Enter the Reset Code", "Please enter your Reset Code sent to your phone number:")
                    if (int(ResetCodeCheck) == RecoveryCode):
                        Reset = True
                        NewPassword = askstring("Success", "Enter your new password")
                        mycursor = mydb.cursor()
                        #updating the password of a specific account
                        UpdateSQL = "UPDATE usercredentials SET Password = %s WHERE Username = %s"
                        passworddata = (NewPassword), (ResetLoginUsername1)
                        mycursor.execute(UpdateSQL, passworddata)
                        mydb.commit()
                        self.show_frame(LoginFrame)
                    else:
                        msg.showinfo("Error","Invalid Reset Code")
                        return
        if Reset == False:
            msg.showinfo("Error","Invalid Username")
    
    def haversine(self,lon1, lat1, lon2, lat2):
        """ Calculate the distance between two points on the earth (specified in decimal degrees) """
        #convert decimal degrees to radians
        lon1, lat1, lon2, lat2 = map(radians, [lon1, lat1, lon2, lat2])
        #haversine formula
        dlon = lon2 - lon1
        dlat = lat2 - lat1
        a = sin(dlat / 2) ** 2 + cos(lat1) * cos(lat2) * sin(dlon / 2) ** 2
        c = 2 * asin(sqrt(a))
        # Radius of earth in kilometers is 6371
        km = 6371 * c
        km = round(km, 3)
        return km


    def makeMap(self,data):
                INST_LAT = 54.597045
                INST_LONG = -5.936478
                m = folium.Map(
                        location=[INST_LAT, INST_LONG],
                        zoom_start=9,
                        tiles='Stamen Terrain'
                )
                tooltip = 'Click me!'
                # make INST Marker
                folium.Marker([INST_LAT, INST_LONG], popup='<i>Your Location</i>', tooltip=tooltip).add_to(m)
                # make Users Marker
                folium.Marker([data["result"]["latitude"],data["result"]["longitude"]], popup='<i>User Input</i>',
                                          tooltip=tooltip).add_to(m)
                # list of tuples containing lat and longs to create line
                points = [(INST_LAT, INST_LONG), (data["result"]["latitude"], data["result"]["longitude"])]
                # calculate distance between two points using Haversine Formula. Uses Km
                distance = self.haversine(data["result"]["longitude"], data["result"]["latitude"], INST_LONG, INST_LAT)
                folium.PolyLine(points, color="red", weight=2.5, opacity=1, tooltip=str(distance) + "km").add_to(m)
                m.save('ElectricianMaps.html')
                # Auto open
                webbrowser.open('file://' + os.path.realpath('ElectricianMaps.html'))
                            
    
    # Obtain the dictionary of data associated with postcode.
    def getLatLong(self,inputtedPostcode):
        if " " in inputtedPostcode:
            msg.showinfo("Error","Please enter the postcode again without a space")
            return
        mustContain4 = "BT"
       #Checking to ensure the email includes an @
        if mustContain4 in inputtedPostcode:
            print ("Postcode Validated")
        else:
            msg.showinfo("Error","Invalid Postcode")
            return
            
            
        try:
            ssl._create_default_https_context = ssl._create_unverified_context
            res = urllib.request.urlopen("http://api.postcodes.io/postcodes/{}".format(inputtedPostcode)).read()
            data = json.loads(res)
            self.makeMap(data)
        
        except Exception as e:
            print(e)
            msg.showerror("Error!", e)

    def sendEmail(self,subject, msg, recipientemail,text):
       # try:
            server = smtplib.SMTP('smtp.gmail.com',587)
            server.starttls()
            emailAddress = "darsonemailservice@gmail.com"
            password = ""
            server.login(emailAddress, password)
            print(msg)
            server.sendmail(emailAddress, recipientemail,text)
            #server.quit()
            print("Email Sent")
    
    #Generetaing Reports
    def YearlyReport(self):
            YearlyReportMonths = ('January', 'Feburary', 'March', 'April', 'May', 'June', 'July', 'August'
               ,'September','October','November','December')
            y_pos = np.arange(len(YearlyReportMonths))
            #Getting all the data from jobs table
            Jan1 = Func.Jan
            Feb1 = Func.Feb
            Mar1 = Func.Mar 
            Apr1 = Func.Apr 
            May1 = Func.May
            June1 = Func.June 
            July1 = Func.July 
            Aug1 = Func.Aug 
            Sep1 = Func.Sep 
            Oct1 = Func.Oct 
            Nov1 = Func.Nov 
            Dec1 = Func.Dec
            #Data variable
            performance = [Jan1,Feb1,Mar1,Apr1,May1,June1,July1,Aug1,Sep1,Oct1,Nov1,Dec1]
            plt.bar(y_pos, performance, align='center', alpha=0.5)
            plt.xticks(y_pos, YearlyReportMonths)
            #Setting the label and title
            plt.ylabel('Jobs Completed')
            plt.title('Jobs Completed Per Month')
            #email address which emails are sent from
            emailAddress = "darsonemailservice@gmail.com"
            #Executing command to get data from the databse
            mycursor.execute("SELECT managerID,managerEmailAddress FROM managerContact")
            myresult = mycursor.fetchall()
            print (myresult)
            #Displaying the subject of the email
            messagebox.showinfo("Emailed","This report has also been sent to your email.")
            subject = "Generated Reports"
            recipientemail = myresult[0][1]
            msg = MIMEMultipart()
            msg['From']  = emailAddress
            msg['To'] = recipientemail
            msg['Subject'] = subject
            body = "Here is the generated report."
            msg.attach(MIMEText(body,'plain'))
            #Attaching the image to the email
            filename = "WeeklyReport.png"
            attachment = open(filename,'rb')
            part = MIMEBase('application','octet-stream')
            part.set_payload((attachment).read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition',"attachment; filename= "+filename)
            msg.attach(part)
            text = msg.as_string()
            #Running the command to send the email
            self.sendEmail(subject,msg,recipientemail,text)                
            #Showing the graph
            plt.show()
    
    #Generating Client Feedback Report
    def ClientFeedbackReport(self):
            FeedbackSections = ('Time','Friendliness','System','Cost','Overall')
            y_pos = np.arange(len(FeedbackSections))
            P1 = 0
            P2 = 0
            P3 = 0
            P4 = 0
            P5 = 0
            #Getting the data from the clientFeedback table
            mycursor.execute("SELECT SystemFeedback,timeFeedback,FriendlinessFeedback,CostFeedback,OverallFeedback FROM clientFeedback")
            myresults = mycursor.fetchall()
            #Looping through it
            for row in myresults:
                #Adding the results
                P1 += int(row[0])
                P2 += int(row[1])
                P3 += int(row[2])
                P4 += int(row[3]) 
                P5 += int(row[4]) 

            #Checking how many individual sets of feedback there were
            Length = len(myresults)
            #Calculating average
            P1Total = P1 / Length
            P2Total = P2 / Length
            P3Total = P3 / Length
            P4Total = P4 / Length
            P5Total = P5 / Length
            performance = [P1Total,P2Total,P3Total,P4Total,P5Total]
            plt.bar(y_pos, performance, align='center', alpha=0.5)
            plt.xticks(y_pos, FeedbackSections)
            plt.ylabel('Average Rating (1-5)')
            plt.title('Feedback Sections')
            messagebox.showinfo("Emailed","This report has also been sent to your email.")
            #email address which emails are sent from
            emailAddress = "darsonemailservice@gmail.com"
            #Executing command to get data from the databse
            mycursor.execute("SELECT managerID,managerEmailAddress FROM managerContact")
            myresult = mycursor.fetchall()
            print (myresult)
            #Displaying the subject of the email
            subject = "Generated Reports"
            recipientemail = myresult[0][1]
            msg = MIMEMultipart()
            msg['From']  = emailAddress
            msg['To'] = recipientemail
            msg['Subject'] = subject
            body = "Here is the generated report."
            msg.attach(MIMEText(body,'plain'))
            #Attaching the image to the email
            filename = "CustomerFeedback.png"
            attachment = open(filename,'rb')
            part = MIMEBase('application','octet-stream')
            part.set_payload((attachment).read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition',"attachment; filename= "+filename)
            msg.attach(part)
            text = msg.as_string()
            #Running the command to send the email
            self.sendEmail(subject,msg,recipientemail,text)                
            #Showing the graph
            plt.show()

    #InvoiceGenerationFunction
    def invoiceGeneration(self,EmployView):
        items = EmployView.selection()
        print(items)
        currentjobData = []
        for i in items:
            currentjobData.append(EmployView.item(i)['values'])
        print(currentjobData[0][0])
        print(currentjobData[0][5])
        #Checking if the job is complete
        if (str(currentjobData[0][5]) != "Complete"):
            #Outputting an error if nots completed.
            msg.showinfo("Error","Job is not complete.")
            return
        InsertSQL = "INSERT INTO pastJobs(firstname,secondname,gender,jobtype,PostCodeOfWork,jobStatus) VALUES(%s,%s,%s,%s,%s,%s)"
        createaccountdata = (currentjobData[0][0],currentjobData[0][1],currentjobData[0][2],currentjobData[0][3],currentjobData[0][4],currentjobData[0][5])
        mycursor.execute(InsertSQL, createaccountdata)
        mydb.commit()
        #Asking the user how many hours it took for the job to be completed.
        HourlyWage = 7.50
        TotalHours = askstring("Total Hours","Please input the total hours to complete the job")
        #Calculating the Cost
        TotalCost = float(HourlyWage) * int(TotalHours)
        print(TotalCost)
        #Creating the dcoument
        mydoc = docx.Document("Invoice.docx")
        #Adding details to the document
        mydoc.add_paragraph("Invoice for recent requested work")
        mydoc.add_paragraph("Name: "+str(currentjobData[0][0])+" "+str(currentjobData[0][1]))
        mydoc.add_paragraph("Your recent work of "+str(currentjobData[0][3])+" has now been completed and the following invoice has been generated. \nPlease make it payable to Darson Services. \nFeel free to leave feedback next time you login to the system")
        mydoc.add_paragraph("Subtotal: £"+str(TotalCost))
        #Calculating Tax
        Tax = TotalCost * 0.20
        mydoc.add_paragraph("Tax: "+str(Tax))
        Total = TotalCost + Tax
        mydoc.add_paragraph("Total: "+str(Total))
        #Saving the document
        mydoc.save("Invoice.docx")
        #Informing the user the document was successfully created.
        msg.showinfo("Success","Document Successfully created")
        UpdateSQL = "DELETE From activeJobs WHERE jobStatus = 'Complete'"
        mycursor.execute(UpdateSQL)
        mydb.commit()
    def managerCreateTreeview(self,EmployView):
            # Delete Everything From Treeview
        Remove = EmployView.get_children()
        for child in Remove:
                    EmployView.delete(child)
            # SELECT data from table
        mycursor.execute("SELECT firstname,secondname,gender,jobtype,PostcodeOfWork,jobStatus FROM activeJobs")
        myresults = mycursor.fetchall()
        #Insert obtained data into treeview widget
        for i in myresults:
                    EmployView.insert("", "end", text="", values=(i[0], i[1], i[2], i[3], i[4],i[5]))    

    #Sending Email function 
    def sendEmail(self,subject, msg, recipientemail,text):
       # try:
            server = smtplib.SMTP('smtp.gmail.com',587)
            server.starttls()
            emailAddress = "darsonemailservice@gmail.com"
            password = ""
            server.login(emailAddress, password)
            print(msg)
            server.sendmail(emailAddress, recipientemail,text)
            #server.quit()
            print("Email Sent")
      #  except:
        #    print("Email Not Sent")

        #Attaching the details to the email.  
   # def EmailReports(self):
        
    #Confirming the adding of details       
    def add_customer_details_confirm(self,CreateNewUsername,CreateNewPassword,CreateNewPhoneNumber,CreateNewEmail,CreateNewAddress,CreateNewPostcode,ManagerCreateAge):
       #Getting all the data into variables that the user inputted
       add_customer_username = CreateNewUsername.get()
       add_customer_password = CreateNewPassword.get()
       add_customer_contactnumber = CreateNewPhoneNumber.get()
       add_customer_postcode = CreateNewPostcode.get()
       add_customer_email = CreateNewEmail.get()
       add_customer_address = CreateNewAddress.get()
       add_customer_postcode = CreateNewPostcode.get()
       add_customer_age = ManagerCreateAge.get()
  
       #generating a customerid
       customerid = random.randint(1111, 9999)
       #Carrying out the checks on Length and Presence
       if add_customer_username == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_password == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_contactnumber == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_contactnumber == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_postcode == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_email == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_address == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_postcode == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_age == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       mycursor = mydb.cursor()
       #Getting all the necessary checks for phone numbers, emails and postcodes
       FormatEmail = False
       mustContain = "@"
       mustContain2 = "+"
       mustContain3 = ".com"
       mustContain4 = "BT"
       #Checking to ensure the email includes an @
       if mustContain in add_customer_email:
           print ("Email Validated")    
       else:
           #Outputting error
           msg.showinfo("Invalid Email","You have entered an invalid email")
           return
       #Checking to ensure 
       if mustContain2 in add_customer_contactnumber:
           print ("Phone Number Validated") 
       else:
           msg.showinfo("Invalid Phone Number","Please follow the correct format of +44")
           return
       #Ensuring the email includes a .com in it 
       if mustContain3 in add_customer_email:
           print ("Email Validated 2")  
       else:
           #Outputting an error to the user to inform them
           msg.showinfo("Invalid Email","You have entered an invalid email")
           return
       #Ensuring the postcode includes 'BT' in it 
       if mustContain4 in add_customer_postcode:
           print ("Postcode Validated") 
       else:
           msg.showinfo("Invalid Postcode","You have entered an invalid Postcode")
           return 

            
       #Inserting some data into the client info table
       InsertSQL = "INSERT INTO clientInfo(clientID,firstname,age,address,postcode) VALUES(%s,%s,%s,%s,%s)"
       addcustomerdata = (str(customerid), str(add_customer_username),str(add_customer_age),str(add_customer_address), str(add_customer_postcode))
       mycursor.execute(InsertSQL, addcustomerdata)
       mydb.commit()
       #Inserting some data into the clientContact table
       InsertSQL2 = "INSERT INTO clientContact(clientID,clientPhoneNumber,clientEmailAddress) VALUES(%s,%s,%s)"
       addcustomerdata2 = (customerid, str(add_customer_contactnumber),str(add_customer_email))
       mycursor.execute(InsertSQL2, addcustomerdata2)
       mydb.commit()
       #Inserting the login details to the user credentials table.
       InsertSQL3 = "INSERT INTO usercredentials(userID,username,password,stakeholder) VALUES(%s,%s,%s,%s)"
       addcustomerdata3 = (customerid, str(add_customer_username),str(add_customer_password),str("Customer"))
       mycursor.execute(InsertSQL3, addcustomerdata3)
       mydb.commit()
       #Outputting a confirmation message.
       msg.showinfo(title="Confirmation", message="Added. You have been sent a confimration text message")
       #Sending Text message.
       client.publish(
                PhoneNumber=str(add_customer_contactnumber),
                Message=str("Account successfully created")
                )
       #Resetting all the fields.
       CreateNewUsername.set("")
       CreateNewPassword.set("")
       CreateNewAddress.set("")
       CreateNewPhoneNumber.set("")
       CreateNewPostcode.set("")
       CreateNewEmail.set("")




    def add_customer_details_confirm1(self,CreateNewUsername,CreateNewPassword,CreateNewPhoneNumber,CreateNewEmail,CreateNewAddress,CreateNewPostcode,ManagerCreateAge,ManagerCreateAccountLoginType):
       #Getting all the data into variables that the user inputted
       add_customer_username = CreateNewUsername.get()
       add_customer_password = CreateNewPassword.get()
       add_customer_contactnumber = CreateNewPhoneNumber.get()
       add_customer_postcode = CreateNewPostcode.get()
       add_customer_email = CreateNewEmail.get()
       add_customer_address = CreateNewAddress.get()
       add_customer_postcode = CreateNewPostcode.get()
       add_customer_age = ManagerCreateAge.get()
       add_customer_logintype = ManagerCreateAccountLoginType.get()
  
       #generating a customerid
       customerid = random.randint(1111, 9999)
       #Carrying out the checks on Length and Presence
       if add_customer_username == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_password == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_contactnumber == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_contactnumber == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_postcode == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_email == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_address == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_postcode == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       if add_customer_age == "":
           msg.showinfo("Error","Failed to complete fields")
           return
       mycursor = mydb.cursor()
       #Getting all the necessary checks for phone numbers, emails and postcodes
       FormatEmail = False
       mustContain = "@"
       mustContain2 = "+"
       mustContain3 = ".com"
       mustContain4 = "BT"
       #Checking to ensure the email includes an @
       if mustContain in add_customer_email:
           print ("Email Validated")    
       else:
           #Outputting error
           msg.showinfo("Invalid Email","You have entered an invalid email")
           return
       #Checking to ensure 
       if mustContain2 in add_customer_contactnumber:
           print ("Phone Number Validated") 
       else:
           msg.showinfo("Invalid Phone Number","Please follow the correct format of +44")
           return
       #Ensuring the email includes a .com in it 
       if mustContain3 in add_customer_email:
           print ("Email Validated 2")  
       else:
           #Outputting an error to the user to inform them
           msg.showinfo("Invalid Email","You have entered an invalid email")
           return
       #Ensuring the postcode includes 'BT' in it 
       if mustContain4 in add_customer_postcode:
           print ("Postcode Validated") 
       else:
           msg.showinfo("Invalid Postcode","You have entered an invalid Postcode")
           return 

       if add_customer_logintype == "Customer":     
           #Inserting some data into the client info table
           InsertSQL = "INSERT INTO clientInfo(clientID,firstname,age,address,postcode) VALUES(%s,%s,%s,%s,%s)"
           addcustomerdata = (str(customerid), str(add_customer_username),str(add_customer_age),str(add_customer_address), str(add_customer_postcode))
           mycursor.execute(InsertSQL, addcustomerdata)
           mydb.commit()
           #Inserting some data into the clientContact table
           InsertSQL2 = "INSERT INTO clientContact(clientID,clientPhoneNumber,clientEmailAddress) VALUES(%s,%s,%s)"
           addcustomerdata2 = (customerid, str(add_customer_contactnumber),str(add_customer_email))
           mycursor.execute(InsertSQL2, addcustomerdata2)
           mydb.commit()
           #Inserting the login details to the user credentials table.
           InsertSQL3 = "INSERT INTO usercredentials(userID,username,password,stakeholder) VALUES(%s,%s,%s,%s)"
           addcustomerdata3 = (customerid, str(add_customer_username),str(add_customer_password),str("Customer"))
           mycursor.execute(InsertSQL3, addcustomerdata3)
           mydb.commit()
           #Outputting a confirmation message.
           msg.showinfo(title="Confirmation", message="Added. You have been sent a confimration text message")
           #Sending Text message.
           client.publish(
                PhoneNumber=str(add_customer_contactnumber),
                Message=str("Account successfully created")
                )
           #Resetting all the fields.
           CreateNewUsername.set("")
           CreateNewPassword.set("")
           CreateNewAddress.set("")
           CreateNewPhoneNumber.set("")
           CreateNewPostcode.set("")
           CreateNewEmail.set("")

       if add_customer_logintype == "Electrician":     
           #Inserting some data into the client info table
           InsertSQL = "INSERT INTO electricianInfo(electricianID,electricianName,electricianAddress) VALUES(%s,%s,%s)"
           addcustomerdata = (str(customerid), str(add_customer_username),str(add_customer_address))
           mycursor.execute(InsertSQL, addcustomerdata)
           print("Complete 1")
           mydb.commit()
           #Inserting some data into the clientContact table
           InsertSQL2 = "INSERT INTO electricianContact(electricianID,electricianPhoneNumber,electricianEmail) VALUES(%s,%s,%s)"
           addcustomerdata2 = (customerid, str(add_customer_contactnumber),str(add_customer_email))
           mycursor.execute(InsertSQL2, addcustomerdata2)
           print("Complete 2")
           mydb.commit()
           #Inserting the login details to the user credentials table.
           InsertSQL3 = "INSERT INTO usercredentials(userID,username,password,stakeholder) VALUES(%s,%s,%s,%s)"
           addcustomerdata3 = (customerid, str(add_customer_username),str(add_customer_password),str("Electrician"))
           mycursor.execute(InsertSQL3, addcustomerdata3)
           mydb.commit()
           #Outputting a confirmation message.
           msg.showinfo(title="Confirmation", message="Added. You have been sent a confimration text message")
           #Sending Text message.
           client.publish(
                PhoneNumber=str(add_customer_contactnumber),
                Message=str("Account successfully created")
                )
           #Resetting all the fields.
           CreateNewUsername.set("")
           CreateNewPassword.set("")
           CreateNewAddress.set("")
           CreateNewPhoneNumber.set("")
           CreateNewPostcode.set("")
           CreateNewEmail.set("")
       if add_customer_logintype == "Manager":     
           #Inserting some data into the client info table
           InsertSQL = "INSERT INTO managerInfo(managerID,managerFirstname,DateOfBirth,managerGender) VALUES(%s,%s,%s,%s)"
           addcustomerdata = (str(customerid), str(add_customer_username),str(add_customer_age), str(add_customer_postcode))
           mycursor.execute(InsertSQL, addcustomerdata)
           mydb.commit()
           #Inserting some data into the clientContact table
           InsertSQL2 = "INSERT INTO managerContact(managerID,managerPhoneNumber,managerEmailAddress) VALUES(%s,%s,%s)"
           addcustomerdata2 = (customerid, str(add_customer_contactnumber),str(add_customer_email))
           mycursor.execute(InsertSQL2, addcustomerdata2)
           mydb.commit()
           #Inserting the login details to the user credentials table.
           InsertSQL3 = "INSERT INTO usercredentials(userID,username,password,stakeholder) VALUES(%s,%s,%s,%s)"
           addcustomerdata3 = (customerid, str(add_customer_username),str(add_customer_password),str("Manager"))
           mycursor.execute(InsertSQL3, addcustomerdata3)
           mydb.commit()
           #Outputting a confirmation message.
           msg.showinfo(title="Confirmation", message="Added. You have been sent a confimration text message")
           #Sending Text message.
           client.publish(
                PhoneNumber=str(add_customer_contactnumber),
                Message=str("Account successfully created")
                )
           #Resetting all the fields.
           CreateNewUsername.set("")
           CreateNewPassword.set("")
           CreateNewAddress.set("")
           CreateNewPhoneNumber.set("")
           CreateNewPostcode.set("")
           CreateNewEmail.set("")

    #Creating manager announcement
    def Manager_Announcement(self,AnnouncementTitle,AnnouncementDetails,UserGroupType,UserGroupType2):
        #Getting the announcement title and details from what the user 
        AnnouncementTitle1 =  AnnouncementTitle.get()
        AnnouncementDetails1 = AnnouncementDetails.get()
        #Getting the usegroups
        UserGroupType1 = UserGroupType.get()
        UserGroupType3 = UserGroupType2.get()
        #Checkin if both user groups are in use
        if (UserGroupType3 == ""):
            #Print message for debugging
            print("2nd User Group not used")
        #Presence check for both user groups
        if (UserGroupType1 == "" and UserGroupType3 == ""):
            #Displaying unsuccessful message
            msg.showinfo("Invalid User Group selected")
            return
        #Presence check for title and details
        if (AnnouncementTitle1 == "" or AnnouncementDetails1 == ""):
            #Displaying unsuccessful message
            msg.showinfo("Error","Invalid details entered")
            return
        #Checking the user group
        if (UserGroupType1 == "Customer"):
            #Getting details from the table
            mycursor.execute("SELECT clientID,clientPhoneNumber FROM clientContact")
            myresult = mycursor.fetchall()
            #Looping through the results
            for i in myresult:
                #Sending text messages to the users.
                client.publish(PhoneNumber = str(myresult[1],message = str(AnnouncementTitle1)+"   "+str(AnnouncementDetails1)))
            #Displaying success message to the user.
            msg.showinfo("Success","Announcement has been published")
        #Checking the user group
        if (UserGroupType1 == "Manager"):
            #Getting the details about the user group
            mycursor.execute("SELECT managerID,managerPhoneNumber FROM managerContact")
            myresult = mycursor.fetchall()
            print(myresult)
            #Looping through the collected data
            for i in myresult:
                print(myresult)
                #Getting the phone number
                MessagePhoneNumber = myresult[0][1]
                #Sending text message
                client.publish(PhoneNumber = str(MessagePhoneNumber),Message = str(AnnouncementTitle1)+"   "+str(AnnouncementDetails1))
            #Displaying success message
            msg.showinfo("Success","Announcement has been published")
            #Resetting the variables
            AnnouncementTitle.set("")
            AnnouncementDetails.set("")
        if (UserGroupType1 == "Electrician"):
            #Getting the details about the user group
            mycursor.execute("SELECT electricianID,electricianPhoneNumber FROM electricianContact")
            myresult = mycursor.fetchall()
            print(myresult)
            #Looping through the collected data
            for i in myresult:
                print(myresult)
                #Getting the phone number
                MessagePhoneNumber = myresult[0][1]
                #Sending text message
                client.publish(PhoneNumber = str(MessagePhoneNumber),Message = str(AnnouncementTitle1)+"   "+str(AnnouncementDetails1))
            #Displaying success message
            msg.showinfo("Success","Announcement has been published")
            #Resetting the variables
            AnnouncementTitle.set("")
            AnnouncementDetails.set("")

    #Submit stock function (allowing the adding of stock to the table)
    def SubmitStock(self,AmountOfStockAdded,StockBeingAdded):
        #Message for debugging
        print(AmountOfStockAdded)
        #Getting the variables from the fields the user inputted
        AmountOfStockAdded1 = AmountOfStockAdded.get()
        StockBeingAdded1 = StockBeingAdded.get()
        #Presence Check for the variables inputted by the user
        if (AmountOfStockAdded1 == "" or StockBeingAdded1 == ""):
            #Display unsuccessful message to the user.
            msg.showinfo("Error","Invalid data entered")
            return
        #Getting data from the database
        mycursor.execute("SELECT stockType,currentAmount FROM currentStock")
        myresult = mycursor.fetchall()
        Added = False
        #Looping through the results
        for row in myresult:
            print(myresult)
            print(StockBeingAdded1)
            #Checking if the current row is the stock that is being added
            if (row[0] == StockBeingAdded1):
                #Getting the new total amount of stock
                NewTotal = int(row[1]) + int(AmountOfStockAdded1)
                print(NewTotal)
                #Updating the new total in the SQL database
                UpdateSQL = "UPDATE currentStock SET currentAmount = %s WHERE stockType = %s"
                updatestockdata = (str(NewTotal),str(StockBeingAdded1))
                mycursor.execute(UpdateSQL, updatestockdata)
                mydb.commit()
                Added = True
                msg.showinfo("Success","Stock successfully added.")
                break
            
        if Added == False:
            UpdateSQL = "INSERT INTO currentStock(stockType,currentAmount) VALUES(%s,%s)"
            updatestockdata = (str(StockBeingAdded1),str(AmountOfStockAdded1))
            mycursor.execute(UpdateSQL, updatestockdata)
            mydb.commit()
            msg.showinfo("Success","Stock successfully added.")
                
                

    #Edit stock amount function           
    def EditStockAmount(self,EmployView):
        #Getting the item selected by the user
        items = EmployView.selection()
        currentjobData = []
        for i in items:
            currentjobData.append(EmployView.item(i)['values'])
        print(currentjobData[0][0])
        print(currentjobData[0][1])
        #Asking the user how much they would like to remove
        StockAmountToRemove = askstring("Reduce Stock Levels","Select how much stock you would like to remove")
        print(StockAmountToRemove)
        #getting data from the database
        mycursor.execute("SELECT stockType,currentAmount FROM currentStock")
        myresult = mycursor.fetchall()
        #Looping through the data recieved.
        for row in myresult:
            print(row[0])
            #Checking if the current row is the stock that needs to be edited
            if(row[0] == currentJobData[0][0]):
                #Calculating the new total
                NewTotal = int(currentJobData[0][1]) - StockAmountToRemove
                if (NewTotal == 0):
                    #Displaying success message to the user
                    msg.showinfo("Removed","This stock has been entirely removed")
                    #Updating the SQL database
                    UpdateSQL = "DELETE From currentStock WHERE currentAmount = '0'"
                    mycursor.execute(UpdateSQL, updatestockdata)
                    mydb.commit()
                else:
                    #Display success message to the user
                    msg.showinfo("Removed","This stock has been reduced")
                    #Updating the SQL database
                    UpdateSQL = "UPDATE currentStock SET currentAmount = %s WHERE stockType = %s"
                    updatestockdata = (str(NewTotal),str(currentjobData[0][0]))
                    mycursor.execute(UpdateSQL, updatestockdata)
                    mydb.commit()
                    
            
        

    def managerStockTreeView(self,EmployView):
            # Delete Everything From Treeview
        Remove = EmployView.get_children()
        for child in Remove:
            EmployView.delete(child)
             #SELECT data from table
        mycursor.execute("SELECT stockType,currentAmount FROM currentStock")
        myresults = mycursor.fetchall()
        #Insert obtained data into treeview widget
        for i in myresults:
                    EmployView.insert("", "end", text="", values=(i[0], i[1]))     
            
    def treeviewEmployeeUpdate3(self,EmployView,UserGroupTypeCheck):
        # Delete Everything From Treeview
        Remove = EmployView.get_children()
        print("Check1")
        for child in Remove:
            EmployView.delete(child)
             #SELECT data from table
        print("Check2")
        if (UserGroupTypeCheck == "Customer"):
            mycursor.execute("SELECT userID,username,password,stakeholder FROM userCredentials")
            myresults = mycursor.fetchall()
            #Insert obtained data into treeview widget
            print(myresults)
            for i in myresults:
                if (i[3] == "Customer"):
                    EmployView.insert("", "end", text="", values=(i[0], i[1], i[2], i[3]))
                        
        if (UserGroupTypeCheck == "Electrician"):
            mycursor.execute("SELECT userID,username,password,stakeholder FROM userCredentials")
            myresults = mycursor.fetchall()
            #Insert obtained data into treeview widget
            for i in myresults:
                if (i[3] == "Electrician"):
                    EmployView.insert("", "end", text="", values=(i[0], i[1], i[2], i[3]))
                        
        if (UserGroupTypeCheck == "Manager"):
            mycursor.execute("SELECT userID,username,password,stakeholder FROM userCredentials")
            myresults = mycursor.fetchall()
            print(myresults)
            #Insert obtained data into treeview widget
            for i in myresults:
                if (i[3] == "manager"):
                    EmployView.insert("", "end", text="", values=(i[0], i[1], i[2], i[3]))

    #Creating tree view
    def CreateTreeView3(self, UserGroupType):
        #Getting the inputted data from the user
        UserGroupTypeCheck = UserGroupType.get()
        #Presence Check
        if (UserGroupTypeCheck == ""):
            #Outputting error 
            msg.showinfo("Error","Invalid User Group Selected")
            return
        #Creating all the heading and components of the tree view
        EmployView=ttk.Treeview(self)
        EmployView['columns']=("UserID","Username","Password","Stakeholder")
        EmployView.place(x=100,y=300)
        EmployView.heading("#0",text="",anchor="w")
        EmployView.column("#0",anchor="center",width=5,stretch=tk.NO)
        EmployView.heading("UserID",text="User ID",anchor="w")
        EmployView.column("UserID",anchor="center",width=100)
        EmployView.heading("Username",text="Username",anchor="w")
        EmployView.column("Username",anchor="center",width=100)
        EmployView.heading("Password",text="Username",anchor="w")
        EmployView.column("Password",anchor="center",width=100)
        EmployView.heading("Stakeholder",text="Stakeholder",anchor="w")
        EmployView.column("Stakeholder",anchor="center",width=100)
        #Bind enter to allow the user to select
        EmployView.bind("<Return>",lambda j:self.managerDeleteAccount(EmployView))
        #Creating Scroll Var
        EmployViewScrollbar=ttk.Scrollbar(self,orient="vertical",command=EmployView.yview)
        EmployView.configure(yscroll=EmployViewScrollbar.set)
        self.treeviewEmployeeUpdate3(EmployView,UserGroupTypeCheck)

    #Edit stock function
    def EditStockAmount(self,EmployView):
        #Getting the items the user selected
        items = EmployView.selection()
        currentjobData = []
        #Loopinf through the items
        for i in items:
            currentjobData.append(EmployView.item(i)['values'])
        print(currentjobData[0][0])
        print(currentjobData[0][1])
        #Asking the user how much they would like to remove
        StockAmountToRemove = askstring("Reduce Stock Levels","Select how much stock you would like to remove")
        print(StockAmountToRemove)
        #Getting data from the database
        mycursor.execute("SELECT stockType,currentAmount FROM currentStock")
        myresult = mycursor.fetchall()
        for row in myresult:
            print(row[0])
            #Checking if the current row is the one selected by the user
            if(row[0] == currentjobData[0][0]):
                #Calculating new total
                NewTotal = int(currentjobData[0][1]) - int(StockAmountToRemove)
                #Checking if it has run out
                if (NewTotal == 0):
                    #Displaying success message to the user
                    msg.showinfo("Removed","This stock has been entirely removed")
                    #Updating database
                    UpdateSQL = "UPDATE currentStock SET currentAmount = %s WHERE stockType = %s"
                    updatestockdata = (str(NewTotal),str(currentjobData[0][0]))
                    mycursor.execute(UpdateSQL, updatestockdata)
                    mydb.commit()
                    UpdateSQL = "DELETE From currentStock WHERE currentAmount = '0'"
                    mycursor.execute(UpdateSQL)
                    mydb.commit()
                else:
                    #Displaying success message to the user
                    msg.showinfo("Removed","This stock has been reduced")
                    #Updating the database
                    UpdateSQL = "UPDATE currentStock SET currentAmount = %s WHERE stockType = %s"
                    updatestockdata = (str(NewTotal),str(currentjobData[0][0]))
                    mycursor.execute(UpdateSQL, updatestockdata)
                    mydb.commit()
        
        
    
    def electricianManageStock(self,EmployView):
            # Delete Everything From Treeview
        Remove = EmployView.get_children()
        for child in Remove:
            EmployView.delete(child)
             #SELECT data from table
        mycursor.execute("SELECT stockType,currentAmount FROM currentStock")
        myresults = mycursor.fetchall()
        #Insert obtained data into treeview widget
        for i in myresults:
                    EmployView.insert("", "end", text="", values=(i[0], i[1]))
                    
    def sendEmail2(self,subject,text,recipientemail):
       server = smtplib.SMTP('smtp.gmail.com',587)
       server.starttls()
       emailAddress = "darsonemailservice@gmail.com"
       password = ""
       server.login(emailAddress, password)
       print(msg)
       server.sendmail(emailAddress,recipientemail,text)
       #server.quit()
       print("Email Sent")
        
    #Updating the job status function
    def UpdateJobStatus(self,EmployView):
        #Getting the job selected by the user
        items = EmployView.selection()
        currentjobData = []
        #Looping through the selection
        for i in items:
            currentjobData.append(EmployView.item(i)['values'])
        #Getting data items from the database
        mycursor.execute("SELECT firstname,secondname FROM activeJobs")
        myresult = mycursor.fetchall()
        for row in myresult:
            if row[0] == currentjobData[0][0] and row[1] == currentjobData[0][1] and currentjobData[0][5] == "Not Started":
                electricianID = askstring("Electrician ID","Please confirm your ElectricianID: ")
                messagebox.showinfo("Selected","This job has been assigned to you and the details sent to your email!")
                UpdateSQL = "UPDATE activeJobs SET electricianID = %s WHERE jobStatus =  %s "
                updatestockdata1 = (str(electricianID),"Not Started")
                mycursor.execute(UpdateSQL,updatestockdata1)
                mydb.commit()
                UpdateSQL = "UPDATE activeJobs SET jobStatus = %s WHERE jobStatus =  %s and firstname = %s"
                updatestockdata2 = ("In Progress","Not Started",currentjobData[0][0])
                mycursor.execute(UpdateSQL,updatestockdata2)
                mydb.commit()
                mycursor.execute("SELECT electricianID,ElectricianEmail FROM electricianContact")
                myresult = mycursor.fetchall()
                recipientemail= ""
              #  for row in myresult:
             #       if row[0] == str(electricianID):
             #           recipientemail=row[1]
             #           print("Check")
             #           print(recipientemail)
                emailAddress = "darsonemailservice@gmail.com"
                subject="Details for your newly accepted job"
                msg = MIMEMultipart()
                msg['From']  = emailAddress
                msg['To'] = recipientemail
                msg['Subject'] = subject
                body = ''
                for line in currentjobData:
                    body += "'%s,%s,%s,%s,%s'"% (line[0],line[1],line[2],line[3],line[4])
                print(body)
                msg.attach(MIMEText(body))
                text = msg.as_string()
                self.sendEmail2(subject,text,recipientemail)
                break
                
            if row[0] == currentjobData[0][0] and row[1] == currentjobData[0][1] and currentjobData[0][5] == "In Progress":
                confirm = messagebox.askquestion ('Complete','Are you sure you want to mark this job as Complete?',icon='warning')
                if confirm == "yes":
                    UpdateSQL = "UPDATE activeJobs SET jobStatus = %s WHERE jobStatus =  %s and firstname = %s"
                    updatestockdata2 = ("Complete","In Progress",currentjobData[0][0])
                    mycursor.execute(UpdateSQL,updatestockdata2)
                    mydb.commit()
                    break


    def electricianActiveJobs(self,EmployView):
            # Delete Everything From Treeview
        Remove = EmployView.get_children()
        for child in Remove:
                    EmployView.delete(child)
            # SELECT data from table
        mycursor.execute("SELECT firstname,secondname,gender,jobtype,PostCodeOfWork,jobStatus FROM activeJobs")
        myresults = mycursor.fetchall()
        #Insert obtained data into treeview widget
        for i in myresults:
                    EmployView.insert("", "end", text="", values=(i[0], i[1], i[2], i[3], i[4],i[5]))    

    
                    



    def customerOrderTreeView(self,EmployView,ConfirmPostCode):
        ConfirmPostCode1 = ConfirmPostCode.get()
        if ConfirmPostCode1 == "":
            msg.showinfo("Error","Failed to complete field")
            return
        if len(ConfirmPostCode1 ) > 10:
            msg.showinfo("Error","Invalid Postcode")
            return
            # Delete Everything From Treeview
        Remove = EmployView.get_children()
        for child in Remove:
                    EmployView.delete(child)
        mycursor.execute("SELECT firstname,secondname,gender,PostcodeOfWork,jobtype,jobStatus FROM activeJobs")
        myresults = mycursor.fetchall()
        Found = False
        for row in myresults:
            print(row[3])
            print(row[5])
            if row[3] == ConfirmPostCode1:
                EmployView.insert("", "end", text="", values=(row[0], row[1], row[2], row[4], row[5]))
                Found = True

        mycursor.execute("SELECT firstname,secondname,gender,PostcodeOfWork,jobtype,jobStatus FROM pastJobs")
        myresults = mycursor.fetchall()
        Found = False
        for row in myresults:
            print(row[3])
            print(row[5])
            if row[3] == ConfirmPostCode1:
                EmployView.insert("", "end", text="", values=(row[0], row[1], row[2], row[4], row[5]))
                Found = True
        if Found == False:
            msg.showinfo("Error","Postcode not found")

            

   # def CreateTreeView6(self):
       

    def ClientRecordFeedback(self,SystemRating,TimeRating,FriendlinessRating,CostRating,OverallRating):
        SystemRating1 = SystemRating.get()
        TimeRating1  = SystemRating.get()
        FriendlinessRating1 = FriendlinessRating.get()
        CostRating1 = CostRating.get()
        OverallRating1 = OverallRating.get()
        if SystemRating1 == "0" and TimeRating1 == "0" and FriendlinessRating1 == "0" and CostRating1 == "0" and OverallRating1 == "0":
            msg.showinfo("Error","Failed to complete the fields")
            return
        InsertSQL = "INSERT INTO clientFeedback(SystemFeedback,timeFeedback,FriendlinessFeedback,CostFeedback,OverallFeedback) VALUES(%s,%s,%s,%s,%s)"
        submitdata = (SystemRating1,TimeRating1,FriendlinessRating1,CostRating1,OverallRating1)
        mycursor.execute(InsertSQL, submitdata)
        mydb.commit()
        msg.showinfo("Thank you!","Your feedback has been recorded.")

    def CustomerBookElectrician(self,CreateNewUsername,CreateNewPassword,JobType1,ManagerCreateAccountLoginType,CreateNewPostcode):
        BookFirstName = CreateNewUsername.get()
        BookSurname = CreateNewPassword.get()
        BookJobType = JobType1.get()
        BookGender = ManagerCreateAccountLoginType.get()
        BookPostCode = CreateNewPostcode.get()

        BookStatus = "Not Started"
        BookElectricianID = "0"
        BookHourlyWage = "7.50"

        if BookFirstName == "":
            msg.showinfo("Error","Please complete the First Name field")
            return
        if BookSurname == "":
            msg.showinfo("Error","Please complete the Surname field")
            return
        if BookJobType == "":
            msg.showinfo("Error","Please complete the Job Type field")
            return
        if BookGender == "":
            msg.showinfo("Error","Please complete the Gender field")
            return
        if BookPostCode == "":
            msg.showinfo("Error","Please complete the Post Code field")
            return
    
        mustContain4 = "BT"
        if mustContain4 in BookPostCode:
           print ("Postcode Validated") 
        else:
           msg.showinfo("Invalid Postcode","You have entered an invalid Postcode")
           return 

        InsertSQL = "INSERT INTO activeJobs(electricianID,firstname,secondname,gender,jobtype,hourlywage,PostCodeOfWork,jobStatus) VALUES(%s,%s,%s,%s,%s,%s,%s,%s)"
        createaccountdata = (BookElectricianID,BookFirstName,BookSurname,BookGender,BookJobType,BookHourlyWage,BookPostCode,BookStatus)
        mycursor.execute(InsertSQL, createaccountdata)
        mydb.commit()
        mycursor.execute("SELECT clientID,postcode FROM clientinfo")
        myresults = mycursor.fetchall()
        Found = False
        ConfirmID = 0
        PhoneNumberToSend = ""
        print(BookPostCode)
        print ("Check1")
        for row in myresults:
            if row[1] == BookPostCode:
                ConfirmID = row[0]
                print("Check 2")
                Found = True
                break

        mycursor.execute("SELECT clientID,clientPhoneNumber FROM clientContact")
        myresults = mycursor.fetchall()
        print("Check 3")
        print (ConfirmID)
        msg.showinfo("Success","Successfully booked an electrician, you will recieve an update shortly")
   
                    
    
class LoginFrame(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background="mint cream")
        Username = StringVar()
        # Variable for the details of Password Entry will be stored
        Password = StringVar()    
        # Title Label
        LoginTitle = tk.Label(self, text="Login", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
                              height=2).pack()
        # Label showing the user where to enter in their username
        UsernameLabel = tk.Label(self, text="Username:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        UsernameEntry = tk.Entry(self, textvariable=Username, width=30, bg="mint cream").pack()
        # Shows the users where to enter in their password
        PasswordLabel = tk.Label(self, text="Password:", bg="mint cream", font="bold, 20").pack()
        PasswordEntry = tk.Entry(self, textvariable=Password, show="*", bg="mint cream", width=30).pack()
        UserGroupLabel = tk.Label(self, text="User Group:", bg="mint cream", font="Bold, 20", width=15, height=1).place(
            x=100, y=200)
        # This variable is for the competitors gender
        LoginType = StringVar()
        # This is a list for the competitors gender
        list1 = ['Customer', 'Electrician', 'Manager']
        # This creates the drop down list
        droplist = OptionMenu(self, LoginType, *list1)
        # This configures the drop down list
        droplist.config(width=30, bg="mint cream")
        # This places the drop down list
        droplist.place(x=275, y=205)
        #Enter button creation which submits the login details entered by the user to be checked.
        Enter = tk.Button(self, text="Enter", font="Bold, 15", bg="aqua", width=17, height=3,
                          command=lambda: controller.Check_Login(Password, Username, LoginType))
        Enter.place(x=100, y=250)
        #Forgot Passowrd button creation which takes the user to forgot password frame
        ForgotPassword = tk.Button(self, text="Forgot Password", font="Bold, 15", bg="aqua", width=17, height=3,
                                   command=lambda: controller.show_frame(ResetPassword))
        ForgotPassword.place(x=100, y=350)
        #Create account button creation which takes the user to the frame to create an account
        CreateAccount = tk.Button(self, text="Create Account", font="Bold, 15", bg="aqua", width=17, height=3,
                                  command=lambda: controller.show_frame(ManagerCreateNewAccount2))
        CreateAccount.place(x=350, y=350)
        #clear account button creation, which clears the data entry fields.
        Clear = tk.Button(self, text="Clear", bg="aqua", command=lambda: controller.Clear(Username,Password),
                          font="Bold, 15", width=17, height=3)
        Clear.place(x=350, y=250)
        #Quit button creation, allowing user to exit the system
        Quit = tk.Button(self, text='Quit', bg="red", command=lambda: controller.quit(), font="Bold, 15", width=30,
                         height=2)
        Quit.place(x=175, y=550)


class MainWindow(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background="mint cream")
        MainTitle = tk.Label(self, text="Main Customer Menu", bg="sky blue", fg="mint cream", font="Bold,  25",
                              width=40, height=2).pack()
        CustomerButton1 = tk.Button(self, text="Book an Electrician", fg="black", font="bold, 20", width=30, height=2,
                                    command=lambda: controller.show_frame(BookElectrician))
        CustomerButton1.config(bg="DeepSkyBlue2")
        CustomerButton1.place(x=125, y=100)
        CustomerButton2 = tk.Button(self, text="See Orders", fg="black", font="bold, 20", width=30, height=2,
                                    command=lambda: controller.show_frame(ViewStatusOfJobs))
        CustomerButton2.place(x=125, y=200)
        CustomerButton2 = tk.Button(self, text="Provide Feedback", fg="black", font="bold, 20", width=30, height=2,
                                    command=lambda: controller.show_frame(CustomerFeedback))
        CustomerButton2.place(x=125, y=300)
        CustomerBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda: controller.show_frame(LoginFrame))
        CustomerBackButton.place(x=150, y=500)
                         


class ManagerWindow(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background="mint cream")
        MainTitle = tk.Label(self, text="Main Manager Menu", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
                             height=2).pack()
        ManagerButton1 = tk.Button(self, text="See Current Jobs", fg="black", font="bold, 20", width=20,
                                   command=lambda: controller.show_frame(ViewCurrentJobs))
        ManagerButton1.config(bg="DeepSkyBlue2")
        ManagerButton1.place(x=50, y=100)
     #   ManagerButton2 = tk.Button(self, text="See Past Jobs", fg="black", font="bold, 20", width=20,
     #                              command=lambda: controller.show_frame(RentBook))
      #  ManagerButton2.place(x=325, y=100)
        ManagerButton3 = tk.Button(self, text="Reports", fg="black", font="bold, 20", width=20,
                                   command=lambda: controller.show_frame(ManagerReports))
        
        ManagerButton3.place(x=325, y=100)
        ManagerButton4 = tk.Button(self, text = "Create New Account", fg = "black", font = 'Bold, 20', width = 20, command = lambda:controller.show_frame(ManagerCreateNewAccount))
        ManagerButton4.place(x=50, y=200)
        ManagerButton6 = tk.Button(self, text = "Make An Announcement", fg = "black", font = 'Bold, 20', width = 20, command = lambda:controller.show_frame(ManagerAccouncementFrame))
        ManagerButton6.place(x=325,y=200)
        ManagerButton7 = tk.Button(self, text = "Manage Stock", fg = "black", font = 'Bold, 20', width = 20, command = lambda:controller.show_frame(ManagerManageStock))
        ManagerButton7.place(x=325,y=300)
        ManagerBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                      command=lambda: controller.MainMenuBack())
        ManagerBackButton.place(x=225, y=500)


class ElectricianWindow(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        inputtedPostcode = StringVar()
        MainTitle = tk.Label(self, text="Electrician Menu", bg="sky blue", fg="mint cream", font="Bold, 25", width=40,
                             height=2).pack()
        
        ElectricianButton1 = tk.Button(self, text="See Current Jobs",fg="black", font="bold, 20", width=30, height=2,
                                       command=lambda: controller.show_frame(ElectricianActiveJobs))
        ElectricianButton1.config(bg="DeepSkyBlue2")
        ElectricianButton1.place(x=125, y=200)
        ElectricianButton2 = tk.Button(self, text="Manage Stock", fg="black", font="bold, 20", width=30, height=2,
                                       command=lambda: controller.show_frame(ElectricianManageStock))
        ElectricianButton2.place(x=125, y=300)
        ElectricianButton3 = tk.Button(self, text="Post Code Lookup", fg="black", font="bold, 20", width=30, height=2,
                                       command=lambda: controller.show_frame(PostCodeLookup))
        ElectricianButton3.place(x=125, y=100)
        ElectricianBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                          command=lambda: controller.MainMenuBack())
        ElectricianBackButton.place(x=200, y=500)


class ResetPassword(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        ResetUsername = StringVar()
        # Variable for the details of Password Entry will be stored
        ResetPassword = StringVar()
        ResetPhoneNumber = StringVar()
        LoginType = StringVar()
        # Title Label
        LoginTitle = tk.Label(self, text="Login", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
                              height=2).pack()
        # Label showing the user where to enter in their username
        UsernameLabel = tk.Label(self, text="Username:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        UsernameEntry = tk.Entry(self, textvariable=ResetUsername, width=30, bg="mint cream").pack()
        PhoneNumberLabel = tk.Label(self, text="Phone Number:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        PhoneNumberEntry = tk.Entry(self, textvariable=ResetPhoneNumber, width=30, bg="mint cream").pack()
        UserGroupLabel = tk.Label(self, text="User Group:", bg="mint cream", font="Bold, 20", width=15, height=1).place(x=100, y=200)
        ResetLoginType = StringVar()
        list1 = ['Customer', 'Electrician', 'Manager']
        droplist = OptionMenu(self, ResetLoginType, *list1)
        # This configures the drop down list
        droplist.config(width=30, bg="mint cream")
        # This places the drop down list
        droplist.place(x=275, y=205)
        Enter = tk.Button(self, text="Enter", font="Bold,20", bg="aqua", width=20,command=lambda: controller.ResetPassword(ResetUsername, ResetLoginType, ResetPhoneNumber))
        Enter.place(x=50, y=250)
        CustomerBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(LoginFrame)).place(x=325,y=250)
class ManagerReports(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        ManagerReportLabel1 = tk.Label(self, text="Manager Reports", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        ManagerReportButton1 = tk.Button(self, text = 'Generate Yearly Report', bg="sky blue",command = lambda:controller.YearlyReport(),font="Bold, 20", width=30,
                                         height=2).place(x=125,y=100)
        ManagerReportButton1 = tk.Button(self, text = 'Generate Client Feedback Report', command=lambda:controller.ClientFeedbackReport(), bg="sky blue",font="Bold, 20", width=30,
                                         height=2).place(x=125,y=200)
        ManagerReportButton2 = tk.Button(self, text = 'Email Reports', command=lambda:controller.EmailReports(), bg="sky blue",font="Bold, 20", width=30,
                                         height=2).place(x=125,y=300)
        CustomerBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(ManagerWindow)).place(x=150,y=550)
class ViewCurrentJobs(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        ManagerCurrentJobsLabel1 = tk.Label(self, text="View Current Jobs", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        UpdateTreeview = tk.Button(self, text = 'Update Tree View', bg="sky blue",command = lambda:controller.managerCreateTreeview,font="Bold, 20", width=20,
                                         height=2).place(x=125,y=100)
        CustomerBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(ManagerWindow)).place(x=150,y=550)

        EmployView=ttk.Treeview(self)
        EmployView['columns']=("firstname","secondname","gender","jobtype","postcode","jobstatus")
        EmployView.place(x=100,y=200)
        EmployView.heading("#0",text="",anchor="w")
        EmployView.column("#0",anchor="center",width=5,stretch=tk.NO)
        EmployView.heading("firstname",text="First Name",anchor="w")
        EmployView.column("firstname",anchor="center",width=80)
        EmployView.heading("secondname",text="Second Name",anchor="w")
        EmployView.column("secondname",anchor="center",width=80)
        EmployView.heading("gender",text="Gender",anchor="w")
        EmployView.column("gender",anchor="center",width=80)
        EmployView.heading("jobtype",text="Job Type",anchor="w")
        EmployView.column("jobtype",anchor="center",width=80)
        EmployView.heading("postcode",text="Postcode",anchor="w")
        EmployView.column("postcode",anchor="center",width=80)
        EmployView.heading("jobstatus",text="Job Status",anchor="w")
        EmployView.column("jobstatus",anchor="center",width=80)
        EmployView.bind("<Return>",lambda j:controller.invoiceGeneration(EmployView))
        controller.managerCreateTreeview(EmployView)



class ManagerCreateNewAccount(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        CreateNewUsername = StringVar()
        CreateNewPassword = StringVar()
        CreateNewPhoneNumber = StringVar()
        CreateNewEmail = StringVar()
        CreateNewAddress = StringVar()
        CreateNewPostcode = StringVar()
        ManagerCreateAge = StringVar()

        ManagerCreateAccountLabel1 = tk.Label(self, text="Create New Account", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
                                       height=2).pack()
        ManagerCreateAccountUsernameLabel = tk.Label(self, text="Fullname:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        ManagerCreateAccountUsernameEntry = tk.Entry(self, textvariable=CreateNewUsername, width=30, bg="mint cream").pack()
        # Shows the users where to enter in their password
        ManagerCreateAccountPasswordLabel = tk.Label(self, text="Password:", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPasswordEntry = tk.Entry(self, textvariable=CreateNewPassword, show="*", bg="mint cream", width=30).pack()
        ManagerCreateAgeLabel = tk.Label(self, text="Age:", bg="mint cream", font="bold, 20").pack()
        #ManagerCreateAge = Spinbox(self, from_ = 1, to = 99, width = 10).pack()
        ManagerCreateAge = Spinbox(self, from_= 0, to = 100, width = 5)
        ManagerCreateAge.pack()
        ManagerCreateAccountUserGroupLabel = tk.Label(self, text="User Group:", bg="mint cream", font="Bold, 20", width=15, height=1).pack()
        # This variable is for the competitors gender
        ManagerCreateAccountLoginType = StringVar()
        # This is a list for the competitors gender
        list2 = ['Customer', 'Electrician', 'Manager']
        # This creates the drop down list
        droplist = OptionMenu(self, ManagerCreateAccountLoginType, *list2)
        # This configures the drop down list
        droplist.config(width=30, bg="mint cream")
        # This places the drop down list
        droplist.pack()
        ManagerCreateAccountEmailLabel = tk.Label(self, text="Email:", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountEmailEntry = tk.Entry(self, textvariable=CreateNewEmail, bg="mint cream", width=30).pack()
        ManagerCreateAccountPhoneNumberLabel = tk.Label(self, text="Phone Number: +44", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPhoneNumberEntry = tk.Entry(self, textvariable=CreateNewPhoneNumber, bg="mint cream", width=30).pack()
        ManagerCreateAccountAddressLabel = tk.Label(self, text="Address", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountAddress = tk.Entry(self, textvariable=CreateNewAddress, bg="mint cream", width=30).pack()
        ManagerCreateAccountPostcodeLabel = tk.Label(self, text="Postcode", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPostcode= tk.Entry(self, textvariable=CreateNewPostcode, bg="mint cream", width=30).pack()
        ManagerCreateAccountSubmit = tk.Button(self, text="Enter", font="Bold, 15", bg="aqua", width=17, height=3,
                                               command=lambda: controller.add_customer_details_confirm1(CreateNewUsername,CreateNewPassword,CreateNewPhoneNumber,CreateNewEmail,CreateNewAddress,CreateNewPostcode,ManagerCreateAge,ManagerCreateAccountLoginType)).pack()
        ManagerCreateReturnButton = tk.Button(self, text = "Back", font="Bold, 15", bg="aqua", width=17, height=3,command=lambda:controller.show_frame(ManagerWindow)).place(x=0,y=550)



class ManagerAccouncementFrame(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        ManagerReportLabel1 = tk.Label(self, text="Make An Announcement", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        AnnouncementTitle = StringVar()
        AnnouncementDetails = StringVar()
        AnnouncementTitleLabel = tk.Label(self, text="Announcement Title:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        AnnouncementTitleEntry = tk.Entry(self, textvariable=AnnouncementTitle, width=30, bg="mint cream").pack()
        AnnouncementDetailsLabel = tk.Label(self, text="Announcement Details:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        AnnouncementDetailsEntry = tk.Entry(self, textvariable=AnnouncementDetails, width=60, bg="mint cream").pack()
        UserGroupLabel = tk.Label(self, text="User Group:", bg="mint cream", font="Bold, 20", width=15, height=1).pack()
        UserGroupType = StringVar()
        list1 = ['Customer', 'Electrician', 'Manager']
        droplist = OptionMenu(self, UserGroupType, *list1)
        droplist.config(width=30, bg="mint cream")
        droplist.pack()
        UserGroupLabel = tk.Label(self, text="Additional User Group:", bg="mint cream", font="Bold, 20", width=15, height=1).pack()
        UserGroupType2 = StringVar()
        list1 = ['Customer', 'Electrician', 'Manager']
        droplist = OptionMenu(self, UserGroupType2, *list1)
        droplist.config(width=30, bg="mint cream")
        droplist.pack()
        Enter = tk.Button(self, text="Enter", font="Bold, 15", bg="aqua", width=20,
                          command=lambda: controller.Manager_Announcement(AnnouncementTitle,AnnouncementDetails,UserGroupType,UserGroupType2))
        Enter.place(x=325,y=500)
        BackButton = tk.Button(self, text="Back", font="Bold,20", width=20,
                                       command=lambda: controller.show_frame(ManagerWindow))
        BackButton.place(x=50, y=500)


class ManagerManageStock(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        AmountOfStockAdded = StringVar()
        StockBeingAdded = StringVar()
        imageTest = PhotoImage(file="RefreshImage.png")
        ManagerLabel1 = tk.Label(self, text="Manage Stock", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        DetailsLabel = tk.Label(self, text="Stock Type:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        DetailsEntry = tk.Entry(self, textvariable=StockBeingAdded, width=30, bg="mint cream").pack()
        AmountToAddLabel = tk.Label(self, text="Stock Ammount:", bg="mint cream", font="Bold, 20", width=15, height=1).pack()
        AmountOfStockAdded = Spinbox(self,from_= 0, to = 100,width = 5)
        AmountOfStockAdded.pack()
        Submit = tk.Button(self, text="Add Stock", font="Bold,  20", width=20, command = lambda:controller.SubmitStock(AmountOfStockAdded,StockBeingAdded)).place(x=50,y=200)
        UpdateTreeview = tk.Button(self, text = "Refresh", font="Bold,  20", width=20,command = lambda:controller.managerStockTreeView(EmployView)).place(x=325,y=200)
        BackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda: controller.show_frame(ManagerWindow))
        BackButton.place(x=50, y=550)

        treeStyle= ttk.Style()
        treeStyle.configure("mystyle.Treeview",highlightthickness=0,bd=0,font=('Georgia',15))
        treeStyle.configure("mystyle.Treeview.Heading",font=('Arial',16,'bold'))
        treeStyle.layout("mystyle.Treeview",[("mystyle.Treeview.treearea",{'sticky':'nswe'})])
                
        EmployView=ttk.Treeview(self)
        EmployView['columns']=("StockType","CurrentStockAvaliable")
        EmployView.place(x=150,y=300)
        EmployView.heading("#0",text="",anchor="w")
        EmployView.column("#0",anchor="center",width=5,stretch=tk.NO)
        EmployView.heading("StockType",text="Stock Type",anchor="w")
        EmployView.column("StockType",anchor="center",width=150)
        EmployView.heading("CurrentStockAvaliable",text="Current Stock Avaliable",anchor="w")
        EmployView.column("CurrentStockAvaliable",anchor="center",width=150)
        EmployViewScrollbar=ttk.Scrollbar(self,orient="vertical",command=EmployView.yview)
        EmployView.configure(yscroll=EmployViewScrollbar.set)
        EmployViewScrollbar.place(x=450,y=300)
        EmployView.bind("<Return>",lambda j:controller.managerStockTreeView(EmployView))
        controller.managerStockTreeView(EmployView)

        

class ViewAccounts(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        UserGroupType = StringVar()
        ManagerDeleteAccount = tk.Label(self, text="View Accounts", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        UserGroupLabel = tk.Label(self, text="Select User Group:", bg="mint cream", font="Bold, 20", width=15, height=1).pack()
        list1 = ['Customer', 'Electrician', 'Manager']
        droplist = OptionMenu(self, UserGroupType, *list1)
        droplist.config(width=30, bg="mint cream")
        droplist.pack()
        DisplayTreeview = tk.Button(self, text = 'View Tree View', bg="sky blue",command = lambda:controller.CreateTreeView3(UserGroupType),font="Bold, 20", width=20,
                                         height=2).place(x=0,y=150)
        UpdateTreeview = tk.Button(self, text = 'Update Tree View', bg="sky blue",command = lambda:controller.CreateTreeView3(UserGroupType),font="Bold, 20", width=20,
                                         height=2).place(x=350,y=150)
        
        CustomerBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(ManagerWindow)).place(x=150,y=550)

class UserCreateNewAccount(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')

 
class ElectricianManageStock(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        AmountOfStockAdded = StringVar()
        StockBeingAdded = StringVar()
        imageTest = PhotoImage(file="RefreshImage.png")
        ManagerLabel1 = tk.Label(self, text="Manage Stock", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        DetailsLabel = tk.Label(self, text="Stock Type:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        DetailsEntry = tk.Entry(self, textvariable=StockBeingAdded, width=30, bg="mint cream").pack()
        AmountToAddLabel = tk.Label(self, text="Stock Ammount:", bg="mint cream", font="Bold, 20", width=15, height=1).pack()
        AmountOfStockAdded = Spinbox(self,from_= 0, to = 100,width = 5)
        AmountOfStockAdded.pack()
        Submit = tk.Button(self, text="Add Stock", font="Bold,  20", width=20, command = lambda:controller.SubmitStock(AmountOfStockAdded,StockBeingAdded)).place(x=50,y=200)
        UpdateTreeview = tk.Button(self, text = "Refresh", font="Bold,  20", width=20,command = lambda:controller.electricianManageStock(EmployView)).place(x=325,y=200)
        BackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda: controller.show_frame(ElectricianWindow))
        BackButton.place(x=50, y=550)

        treeStyle= ttk.Style()
        treeStyle.configure("mystyle.Treeview",highlightthickness=0,bd=0,font=('Georgia',15))
        treeStyle.configure("mystyle.Treeview.Heading",font=('Arial',16,'bold'))
        treeStyle.layout("mystyle.Treeview",[("mystyle.Treeview.treearea",{'sticky':'nswe'})])
                
        EmployView=ttk.Treeview(self)
        EmployView['columns']=("StockType","CurrentStockAvaliable")
        EmployView.place(x=150,y=300)
        EmployView.heading("#0",text="",anchor="w")
        EmployView.column("#0",anchor="center",width=5,stretch=tk.NO)
        EmployView.heading("StockType",text="Stock Type",anchor="w")
        EmployView.column("StockType",anchor="center",width=150)
        EmployView.heading("CurrentStockAvaliable",text="Current Stock Avaliable",anchor="w")
        EmployView.column("CurrentStockAvaliable",anchor="center",width=150)
        EmployViewScrollbar=ttk.Scrollbar(self,orient="vertical",command=EmployView.yview)
        EmployView.configure(yscroll=EmployViewScrollbar.set)
        EmployViewScrollbar.place(x=450,y=300)
        EmployView.bind("<Return>",lambda j:controller.EditStockAmount(EmployView))
        controller.electricianManageStock(EmployView)

       
    

class ElectricianActiveJobs(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        
        ManagerCurrentJobsLabel1 = tk.Label(self, text="View Current Jobs", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        ElectricianBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(ElectricianWindow)).place(x=150,y=550)
        EmployView=ttk.Treeview(self)
        EmployView['columns']=("firstname","secondname","gender","jobtype","postcode","jobstatus")
        EmployView.place(x=100,y=200)
        EmployView.heading("#0",text="",anchor="w")
        EmployView.column("#0",anchor="center",width=5,stretch=tk.NO)
        EmployView.heading("firstname",text="First Name",anchor="w")
        EmployView.column("firstname",anchor="center",width=80)
        EmployView.heading("secondname",text="Second Name",anchor="w")
        EmployView.column("secondname",anchor="center",width=80)
        EmployView.heading("gender",text="Gender",anchor="w")
        EmployView.column("gender",anchor="center",width=80)
        EmployView.heading("jobtype",text="Job Type",anchor="w")
        EmployView.column("jobtype",anchor="center",width=80)
        EmployView.heading("postcode",text="Post Code",anchor="w")
        EmployView.column("postcode",anchor="center",width=80)
        EmployView.heading("jobstatus",text="Job Status",anchor="w")
        EmployView.column("jobstatus",anchor="center",width=80)
        EmployView.bind("<Return>",lambda j:controller.UpdateJobStatus(EmployView))
        controller.electricianActiveJobs(EmployView)
        
        UpdateTreeview = tk.Button(self, text = 'Update Tree View', bg="sky blue",command = lambda:controller.electricianActiveJobs(EmployView),font="Bold, 20", width=20,
                                         height=2).place(x=125,y=100)

class PostCodeLookup(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        inputtedPostcode = StringVar()
        self.configure(background='mint cream')
        ManagerLabel1 = tk.Label(self, text="Post Code Lookup", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        PostcodeSearch = tk.Label(self, text="Enter Postcode:", bg="mint cream", font="bold, 20")
        PostcodeSearch.place(x = 125, y = 100)
        ElectricianPostcodeSearch = tk.Entry(self, textvariable = inputtedPostcode, bg="mint cream", font="bold, 20")
        ElectricianPostcodeSearch.place(x=300,y=100)
        ElectricianPostCodeButton = tk.Button(self, text = "Search", width=20,height = 2,command = lambda:controller.getLatLong(inputtedPostcode.get())).place(x=200,y=200)
        ElectricianBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(ElectricianWindow)).place(x=150,y=550)
class ManagerCreateNewAccount2(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        CreateNewUsername = StringVar()
        CreateNewPassword = StringVar()
        CreateNewPhoneNumber = StringVar()
        CreateNewEmail = StringVar()
        CreateNewAddress = StringVar()
        CreateNewPostcode = StringVar()
        ManagerCreateAge = StringVar()

        ManagerCreateAccountLabel1 = tk.Label(self, text="Create New Account", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
                                       height=2).pack()
        ManagerCreateAccountUsernameLabel = tk.Label(self, text="Fullname:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        ManagerCreateAccountUsernameEntry = tk.Entry(self, textvariable=CreateNewUsername, width=30, bg="mint cream").pack()
        # Shows the users where to enter in their password
        ManagerCreateAccountPasswordLabel = tk.Label(self, text="Password:", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPasswordEntry = tk.Entry(self, textvariable=CreateNewPassword, show="*", bg="mint cream", width=30).pack()
        ManagerCreateAgeLabel = tk.Label(self, text="Age:", bg="mint cream", font="bold, 20").pack()
        #ManagerCreateAge = Spinbox(self, from_ = 1, to = 99, width = 10).pack()
        ManagerCreateAge = Spinbox(self, from_= 0, to = 100, width = 5)
        ManagerCreateAge.pack()
        ManagerCreateAccountEmailLabel = tk.Label(self, text="Email:", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountEmailEntry = tk.Entry(self, textvariable=CreateNewEmail, bg="mint cream", width=30).pack()
        ManagerCreateAccountPhoneNumberLabel = tk.Label(self, text="Phone Number: +44", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPhoneNumberEntry = tk.Entry(self, textvariable=CreateNewPhoneNumber, bg="mint cream", width=30).pack()
        ManagerCreateAccountAddressLabel = tk.Label(self, text="Address", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountAddress = tk.Entry(self, textvariable=CreateNewAddress, bg="mint cream", width=30).pack()
        ManagerCreateAccountPostcodeLabel = tk.Label(self, text="Postcode", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPostcode= tk.Entry(self, textvariable=CreateNewPostcode, bg="mint cream", width=30).pack()
        ManagerCreateAccountSubmit = tk.Button(self, text="Enter", font="Bold, 15", bg="aqua", width=17, height=3,
                                               command=lambda: controller.add_customer_details_confirm(CreateNewUsername,CreateNewPassword,CreateNewPhoneNumber,CreateNewEmail,CreateNewAddress,CreateNewPostcode,ManagerCreateAge)).pack()
        ManagerCreateReturnButton = tk.Button(self, text = "Back", font="Bold, 15", bg="aqua", width=17, height=3,command=lambda:controller.show_frame(LoginFrame)).place(x=0,y=550)

        
        
class ViewStatusOfJobs(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        ConfirmPostCode = StringVar()
        self.configure(background='mint cream')
        ManagerLabel1 = tk.Label(self, text="View Orders", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        DetailsLabel = tk.Label(self, text="Postcode:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        DetailsEntry = tk.Entry(self, textvariable=ConfirmPostCode, width=30, bg="mint cream").pack()
        UpdateTreeview = tk.Button(self, text = 'Update Tree View', bg="sky blue",command = lambda:controller.customerOrderTreeView(EmployView,ConfirmPostCode),font="Bold, 20", width=20,
                                         height=2).place(x=125,y=150)
        ClientBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(MainWindow)).place(x=150,y=550)


        EmployView=ttk.Treeview(self)
        EmployView['columns']=("firstname","secondname","gender","jobtype","jobstatus")
        EmployView.place(x=100,y=200)
        EmployView.heading("#0",text="",anchor="w")
        EmployView.column("#0",anchor="center",width=5,stretch=tk.NO)
        EmployView.heading("firstname",text="First Name",anchor="w")
        EmployView.column("firstname",anchor="center",width=80)
        EmployView.heading("secondname",text="Second Name",anchor="w")
        EmployView.column("secondname",anchor="center",width=80)
        EmployView.heading("gender",text="Gender",anchor="w")
        EmployView.column("gender",anchor="center",width=80)
        EmployView.heading("jobtype",text="Job Type",anchor="w")
        EmployView.column("jobtype",anchor="center",width=80)
        EmployView.heading("jobstatus",text="Job Status",anchor="w")
        EmployView.column("jobstatus",anchor="center",width=80)
        EmployView.bind("<Return>",lambda j:controller.UpdateJobStatus(EmployView))
        
class CustomerFeedback(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.configure(background='mint cream')
        ManagerLabel1 = tk.Label(self, text="Provide Feedback", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()    
        SystemLabel = tk.Label(self, text="System: ", bg="mint cream", font="bold, 20").pack()
        SystemRating = Spinbox(self,from_= 0, to = 5,width = 5)
        SystemRating.pack()
        TimeLabel = tk.Label(self, text="Time: ", bg="mint cream", font="bold, 20").pack()
        TimeRating = Spinbox(self,from_= 0, to = 5,width = 5)
        TimeRating.pack()
        FriendlinessLabel = tk.Label(self, text="Friendliness: ", bg="mint cream", font="bold, 20").pack()
        FriendlinessRating = Spinbox(self,from_= 0, to = 5,width = 5)
        FriendlinessRating.pack()
        CostLabel = tk.Label(self, text="Cost: ", bg="mint cream", font="bold, 20").pack()
        CostRating = Spinbox(self,from_= 0, to = 5,width = 5)
        CostRating.pack()
        OverallLabel = tk.Label(self, text="Overall: ", bg="mint cream", font="bold, 20").pack()
        OverallRating = Spinbox(self,from_= 0, to = 5,width = 5)
        OverallRating.pack()
        ManagerCreateAccountSubmit = tk.Button(self, text="Enter", font="Bold, 15", bg="aqua", width=17, height=3,
                                               command=lambda:controller.ClientRecordFeedback(SystemRating,TimeRating,FriendlinessRating,CostRating,OverallRating)).pack()
        ClientBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(MainWindow)).place(x=150,y=550)  
        

class BookElectrician(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        inputtedPostcode = StringVar()
        self.configure(background='mint cream')
        
        CreateNewUsername = StringVar()
        CreateNewPassword = StringVar()
        CreateNewPhoneNumber = StringVar()
        CreateNewEmail = StringVar()
        CreateNewAddress = StringVar()
        CreateNewPostcode = StringVar()
        ManagerCreateAge = StringVar()
        JobType1 = StringVar()

        ManagerLabel1 = tk.Label(self, text="Book an Electrician", bg="sky blue", fg="mint cream", font="Bold,  25", width=40,
        height=2).pack()
        
        ManagerCreateAccountUsernameLabel = tk.Label(self, text="Firstname:", bg="mint cream", font="bold, 20").pack()
        # Allows the user to enter in their username
        ManagerCreateAccountUsernameEntry = tk.Entry(self, textvariable=CreateNewUsername, width=30, bg="mint cream").pack()
        # Shows the users where to enter in their password
        ManagerCreateAccountPasswordLabel = tk.Label(self, text="Surname:", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPasswordEntry = tk.Entry(self, textvariable=CreateNewPassword, bg="mint cream", width=30).pack()
        ManagerCreateAccountUserGroupLabel = tk.Label(self, text="Gender:", bg="mint cream", font="Bold, 20", width=15, height=1).pack()
        ManagerCreateAccountLoginType = StringVar()
        # This is a list for the competitors gender
        list2 = ['Male', 'Female']
        # This creates the drop down list
        droplist = OptionMenu(self, ManagerCreateAccountLoginType, *list2)
        # This configures the drop down list
        droplist.config(width=30, bg="mint cream")
        # This places the drop down list
        droplist.pack()
        
        ManagerCreateAccountPostcodeLabel = tk.Label(self, text="Postcode", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPostcode= tk.Entry(self, textvariable=CreateNewPostcode, bg="mint cream", width=30).pack()
        ManagerCreateAccountPasswordLabel = tk.Label(self, text="Job Type: ", bg="mint cream", font="bold, 20").pack()
        ManagerCreateAccountPasswordEntry = tk.Entry(self, textvariable=JobType1, bg="mint cream", width=30).pack()
        ManagerCreateAccountSubmit = tk.Button(self, text="Enter", font="Bold, 15", bg="aqua", width=17, height=3,
                                               command=lambda:controller.CustomerBookElectrician(CreateNewUsername,CreateNewPassword,JobType1,ManagerCreateAccountLoginType,CreateNewPostcode)).pack()
        
        ClientBackButton = tk.Button(self, text="Back", fg="black", font="Bold,20", width=20,
                                       command=lambda:controller.show_frame(MainWindow)).place(x=150,y=550)  
        
        
Main = Main()
Main.mainloop() 
